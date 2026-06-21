from pathlib import Path

from docx import Document


SRC = Path(r"D:\Codex\report.docx")
OUT = Path(r"D:\Codex\project\report_updated.docx")


LOGS = [
    ("6月15日", "工作内容：完成 TravelMind AI 选题、业务目标和需求分析，明确平台面向有旅行规划需求的用户，支持当前位置、预算、天数、偏好输入后生成目的地、酒店、美食、打卡点和每日行程。遇到问题：功能较多，容易过度设计。解决方案：按 KISS 原则划分用户、社区、AI、行程、地图与推荐模块。效果：形成清晰的前后端分离项目方案。"),
    ("6月16日", "工作内容：搭建 Spring Boot 3.3 + Java 21 后端工程，配置 Maven、MyBatis Plus、Spring Security、JWT、Redis、Knife4j 等基础能力。遇到问题：登录态、游客访问和受保护功能边界需要统一。解决方案：通过 JWT 过滤器和 SecurityConfig 区分公开接口与登录后功能。效果：完成用户注册、登录、个人中心等基础接口。"),
    ("6月17日", "工作内容：设计并初始化 MySQL 数据库，保留用户、收藏、浏览历史、AI会话、AI消息、游记评论、行程计划等核心表，景点和酒店逐步改为高德联网 POI 数据源。遇到问题：迭代中字段经常新增导致旧库报错。解决方案：增加启动时自动迁移逻辑补齐 summary、parent_id、头像快照等字段。效果：后端重启后可自动适配新表结构。"),
    ("6月18日", "工作内容：搭建 Vue3 + TypeScript + Vite 前端工程，接入 Pinia、Vue Router、Axios、Naive UI、UnoCSS 和 ECharts。遇到问题：页面较多时视觉和交互容易不统一。解决方案：统一 Apple 风格毛玻璃卡片、深浅色主题、按钮动效和路由守卫。效果：首页、景点、酒店、AI、社区、个人中心等页面具备一致体验。"),
    ("6月22日", "工作内容：将景点和酒店推荐调整为高德联网 POI 数据源，支持按城市和关键词在线搜索，并使用真实 POI 信息、图片兜底和价格估算展示列表与详情。遇到问题：本地图片缺失、重复率高，影响浏览体验。解决方案：减少本地静态依赖，优先使用在线搜索结果与图片回退策略。效果：景点和酒店内容更实时，详情页可点击查看。"),
    ("6月23日", "工作内容：集成高德地图 JS SDK 与 Web 服务能力，完成地图标记、当前位置获取、附近酒店推荐、天气查询和路线跳转方案。遇到问题：高德 Key 类型不匹配会出现 USERKEY_PLAT_NOMATCH。解决方案：区分前端 JS Key 与后端 Web 服务 Key，并在环境变量中分别配置。效果：地图、定位、天气和周边推荐可以联动使用。"),
    ("6月24日", "工作内容：重构 AI 页面为聊天式问答，接入 DeepSeek 接口，支持流式输出、上下文记忆、历史会话和超过 20 轮后的会话摘要。遇到问题：长对话传全部历史会消耗 Token 且响应变慢。解决方案：保留近期消息，并把早期内容压缩为 Conversation Summary。效果：AI 能连续追问旅行方案，同时保持响应稳定。"),
    ("6月25日", "工作内容：完善首页推荐、预算示意和 AI 工具页，将预算改为根据当前方案实时计算，并支持车票、酒店、天数、人数、行李清单、最终行程预览等步骤联动。遇到问题：往返车票和酒店费用计算容易叠加错误。解决方案：统一费用算法，车票按单程票价×2×人数，酒店按晚数和房间数计算。效果：预算展示更贴近真实旅行规划。"),
    ("6月26日", "工作内容：完善收藏、浏览历史和个人中心，收藏景点、酒店、游记后以小方块图片卡片展示，重复浏览同一内容时更新到最前。遇到问题：游客状态和不同用户之间的数据容易混淆。解决方案：所有收藏、浏览、行程功能绑定登录用户，游客点击时提示登录。效果：个人主页能按用户独立展示收藏、历史和自己发布的游记。"),
    ("6月29日", "工作内容：完善游记社区，改为图文瀑布流展示，支持发布游记、多图上传、点赞、删除、搜索、详情页评论和评论回复。遇到问题：评论头像和昵称不能手动输入，否则无法和账号信息同步。解决方案：后端通过 token 绑定 userId，渲染时实时读取用户最新头像和昵称，右键删除本人评论。效果：社区评论更接近真实产品体验。"),
    ("6月30日", "工作内容：重构行程计划页面，保存后的行程以纵向卡片展示，进入详情后按 Day 生成横向流程节点，并用 ECharts 展示组成统计。遇到问题：流程节点直接固定为交通、景点、美食、住宿不够准确。解决方案：根据保存行程内容提取景点和美食，酒店从酒店信息中读取，并去掉不必要的交通节点。效果：行程详情更贴合实际保存内容。"),
    ("7月1日", "工作内容：进行项目收尾、联调和部署准备，检查登录、AI问答、天气、地图、POI搜索、酒店推荐、社区评论、收藏历史、行程保存、Git 提交和前后端构建。遇到问题：部署环境中部分 REST 请求和数据库字段与本地不同步。解决方案：统一接口路径，增加自动数据库迁移，并通过 npm build、mvn package 验证。效果：项目达到可演示和继续扩展的状态。"),
]


def set_cell_text_keep_format(cell, text):
    if not cell.paragraphs:
        cell.add_paragraph(text)
        return
    paragraph = cell.paragraphs[0]
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)
    for extra in cell.paragraphs[1:]:
        for run in extra.runs:
            run.text = ""


def main():
    doc = Document(SRC)
    index = 0
    for table in doc.tables:
        for cell in table.rows[2].cells:
            if "专业实践" in cell.text:
                set_cell_text_keep_format(cell, "专业实践II")
        for row_index in (5, 6):
            if index >= len(LOGS):
                break
            date, content = LOGS[index]
            set_cell_text_keep_format(table.cell(row_index, 0), date)
            set_cell_text_keep_format(table.cell(row_index, 1), content)
            index += 1
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
