from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
TEMPLATE = ROOT / "practice-log-template.docx"
OUTPUT = ROOT / "report.docx"

PROJECT_TITLE = "TravelMind AI（AI个性化旅游推荐与智能规划平台）"

LOGS = [
    (
        "6月15日",
        "工作内容：完成项目选题和需求分析，明确平台围绕“用户输入当前位置、预算、天数、偏好后自动生成旅游推荐与行程规划”展开，整理用户、景点、酒店、游记、收藏、浏览历史、AI会话、行程计划等核心模块。\n"
        "遇到问题：最初功能点较多，容易把架构设计得过重，影响开发效率。\n"
        "解决方案：采用 KISS 原则，不做复杂 DDD 分层，确定 Spring Boot + Vue 前后端分离方案，优先保证注册登录、推荐查询、AI问答、收藏历史、游记和行程规划闭环可运行。\n"
        "效果：形成了清晰的项目范围和技术路线，为后续编码打下基础。",
    ),
    (
        "6月16日",
        "工作内容：搭建后端基础工程，创建 Spring Boot 3.3 + Java 21 项目，引入 Spring Security、JWT、MyBatis Plus、MySQL、Redis、Knife4j、Lombok 等依赖，完成统一返回对象、全局异常处理和基础配置。\n"
        "遇到问题：认证接口、普通业务接口和静态资源接口权限容易互相影响。\n"
        "解决方案：在 SecurityConfig 中按路径放行登录注册、接口文档和上传资源，其余接口通过 JWT 校验；封装 JwtUtil 负责令牌生成和解析。\n"
        "效果：后端工程可以正常启动，接口结构清晰，具备继续开发业务模块的基础。",
    ),
    (
        "6月17日",
        "工作内容：设计并初始化数据库表，完成用户、收藏、浏览历史、AI会话、AI消息、游记、评论、行程计划等表结构，使用 MyBatis Plus Mapper 简化 CRUD 开发。\n"
        "遇到问题：早期景点和酒店使用本地静态数据，图片容易失效，数据真实性也难以维护。\n"
        "解决方案：调整方案为高德联网 POI 作为景点和酒店主要数据源，本地数据库只保存用户相关数据、收藏、历史、AI会话和行程等需要持久化的内容。\n"
        "效果：数据库职责更清楚，减少了维护大量景点酒店数据和图片的成本。",
    ),
    (
        "6月18日",
        "工作内容：完成前端工程搭建，使用 Vue3、TypeScript、Vite、Pinia、Vue Router、Axios、Naive UI、UnoCSS 构建基础页面和路由，包括首页、景点、酒店、AI助手、工具、游记、行程、个人中心等。\n"
        "遇到问题：首页按钮、登录页和部分卡片在非悬浮状态下颜色不够清晰，深浅色主题视觉不统一。\n"
        "解决方案：统一液态玻璃卡片、按钮、导航高亮和阴影样式，减少突兀蓝色，改为更贴合旅游主题的渐变和悬浮加深效果。\n"
        "效果：页面风格更接近 Apple iOS 的毛玻璃和圆角卡片体验，深色模式与浅色模式都能正常阅读。",
    ),
    (
        "6月22日",
        "工作内容：实现景点与酒店在线搜索功能，对接高德地图 POI 服务，支持按城市、关键词检索景点和酒店，并在详情页展示名称、地址、经纬度、评分、图片兜底和周边推荐。\n"
        "遇到问题：高德返回的 location 字段有时不是字符串，前端直接 split 会报错；在线图片也可能缺失。\n"
        "解决方案：增加 location 类型判断和坐标解析兜底逻辑；图片展示改为真实 POI 图片优先，缺失时按地点类别使用稳定的在线兜底图。\n"
        "效果：景点和酒店搜索稳定性提升，用户可以点击列表进入详情页查看信息。",
    ),
    (
        "6月23日",
        "工作内容：集成高德地图 JS SDK，在景点详情和地图组件中展示位置标记，并尝试获取用户当前位置，用于后续附近推荐和路线规划。\n"
        "遇到问题：浏览器定位需要 HTTPS 或 localhost 安全环境，并且用户未授权时不会得到当前位置。\n"
        "解决方案：在 AmapView 组件中加入浏览器 Geolocation 判断、授权失败提示和默认城市中心点兜底，保证地图即使无法定位也能打开。\n"
        "效果：地图组件可正常加载，定位失败时页面仍保持可用，不影响景点和酒店浏览。",
    ),
    (
        "6月24日",
        "工作内容：开发 AI 对话页面，将原来的表单式推荐改为类似 ChatGPT/DeepSeek 的问答模式，支持上下文聊天记录、连续追问、旅游规划建议和预算分析。\n"
        "遇到问题：DeepSeek API Key 配置错误或未被后端读取时会返回 401，且在线模型调用成本不可控。\n"
        "解决方案：后端 AiService 同时支持 DeepSeek 和本地 Ollama，默认可使用免费本地模型，DeepSeek 通过环境变量切换；前端展示当前配置并给出错误提示。\n"
        "效果：AI助手既能在线调用，也能在无付费 Key 的情况下使用本地模型完成旅游问答。",
    ),
    (
        "6月25日",
        "工作内容：完善首页推荐和预算示意，根据当前季节展示国内更适合出行的默认推荐；预算图表支持选择不同方案后实时切换金额，并允许手动修改交通、住宿、餐饮、门票等费用。\n"
        "遇到问题：预算图表在 Edge 页面显示不完整，深色模式下中间金额文字对比度不足。\n"
        "解决方案：调整 ECharts 容器高度、响应式布局和图表中心文字颜色，预算数据不再叠加，而是只显示当前选择方案的总额。\n"
        "效果：首页推荐和预算展示更直观，用户可以按方案快速比较旅行成本。",
    ),
    (
        "6月26日",
        "工作内容：实现收藏和浏览历史功能，支持景点、酒店、游记收藏；个人中心改为类似小红书/抖音的小方块图片卡片展示，并记录用户访问过的详情页。\n"
        "遇到问题：同一个景点或酒店多次浏览会重复出现，收藏后的图片和在线列表看到的图片不一致。\n"
        "解决方案：前端建立 favoriteCache 保存在线对象快照，浏览历史按业务类型和目标 ID 去重，重复访问时更新时间并移动到最前。\n"
        "效果：个人中心收藏和浏览记录能正确更新，图片与联网搜索结果保持一致，展示更美观。",
    ),
    (
        "6月29日",
        "工作内容：完善游记社区，支持发布游记、上传多张照片、点赞、收藏、评论、进入详情页查看完整内容，并增加删除游记功能。\n"
        "遇到问题：评论接口和删除接口路径不一致，前端请求 DELETE 后后端提示 Request method not supported 或 No static resource。\n"
        "解决方案：统一 NoteController 的 RESTful 路由，前端 api.ts 对齐后端路径；图片上传使用 multipart 表单，保存成功后回显图片地址。\n"
        "效果：游记发布、评论和删除流程恢复正常，照片数量不再限制，社区模块形成完整闭环。",
    ),
    (
        "6月30日",
        "工作内容：重构行程计划页面，将保存的行程改为纵向单条卡片列表，卡片显示目的地和出行天数；进入详情后展示 AI 整理出的结构化行程和可视化图表。\n"
        "遇到问题：AI 原始文本不适合直接阅读，用户难以快速理解每天的交通、景点、美食和住宿安排。\n"
        "解决方案：前端对行程内容做结构化解析，按单日生成横向时间轴流程图，使用飞机、地铁、公交、古建筑、海滩、山林、碗筷、酒店等统一图标分类。\n"
        "效果：行程详情页更清楚，点击任意节点可以弹窗查看完整建议、避坑贴士和预约方式。",
    ),
    (
        "7月1日",
        "工作内容：完成项目收尾、测试和 Git 管理，检查前后端构建、接口联调、页面交互、图片兜底、AI配置、收藏历史、游记评论、行程可视化等主要功能，并整理 README、Docker、Nginx 和数据库脚本。\n"
        "遇到问题：上传 GitHub 时需要避免提交 .env.local、node_modules、target、dist、日志、上传文件和 IDE 配置，防止泄露 Key 或仓库过大。\n"
        "解决方案：补充 .gitignore，确认敏感文件未进入提交历史，完成 Initial commit，并准备远程仓库推送命令。\n"
        "效果：TravelMind AI 已形成可运行、可演示、可继续扩展的完整项目版本。",
    ),
]


def set_cell_text(cell, text, size=9):
    cell.text = ""
    for i, line in enumerate(text.split("\n")):
        paragraph = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.space_after = Pt(2)
        paragraph.paragraph_format.line_spacing = 1.05
        run = paragraph.add_run(line)
        run.font.size = Pt(size)
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")


def fill_doc():
    doc = Document(str(TEMPLATE))
    log_index = 0

    for table in doc.tables:
        set_cell_text(table.cell(0, 1), PROJECT_TITLE, 10.5)
        for row_idx in (5, 6):
            if log_index >= len(LOGS):
                break
            date_text, log_text = LOGS[log_index]
            set_cell_text(table.cell(row_idx, 0), date_text, 10)
            set_cell_text(table.cell(row_idx, 1), log_text, 8.5)
            log_index += 1

    doc.save(str(OUTPUT))
    print(OUTPUT)


if __name__ == "__main__":
    fill_doc()
