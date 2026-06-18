from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "baogao.docx"


def set_font(run, size=12, bold=False, name="宋体", color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)


def set_para(paragraph, align=None, before=0, after=6, line=1.5, first_line=False):
    if align is not None:
        paragraph.alignment = align
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    if first_line:
        fmt.first_line_indent = Pt(24)


def add_paragraph(doc, text="", size=12, bold=False, align=None, style=None, first_line=False):
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    set_para(p, align=align, first_line=first_line)
    if text:
        run = p.add_run(text)
        set_font(run, size=size, bold=bold)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    set_para(p, before=10 if level == 1 else 6, after=6, line=1.3)
    run = p.add_run(text)
    set_font(run, size=16 if level == 1 else 14, bold=True, name="黑体")
    return p


def set_cell(cell, text, width=None, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT):
    if width:
        cell.width = width
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    cell.text = ""
    for i, line in enumerate(text.split("\n")):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        set_para(p, align=align, after=2, line=1.25)
        r = p.add_run(line)
        set_font(r, size=10.5, bold=bold)


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        set_cell(table.rows[0].cells[i], h, widths[i] if widths else None, True, WD_ALIGN_PARAGRAPH.CENTER)
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            set_cell(cells[i], str(text), widths[i] if widths else None)
    doc.add_paragraph()
    return table


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run("第 ")
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    paragraph._p.append(fld_begin)
    paragraph._p.append(instr)
    paragraph._p.append(fld_end)
    paragraph.add_run(" 页")


def build():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.5)
    add_page_number(section.footer.paragraphs[0])

    styles = doc.styles
    styles["Normal"].font.name = "宋体"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    styles["Normal"].font.size = Pt(12)

    # Cover
    add_paragraph(doc, "专业实践II报告", size=24, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, "", size=12)
    add_paragraph(doc, "题    目：TravelMind AI（AI个性化旅游推荐与智能规划平台）", size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, "课程名称：专业实践II", size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, "班    级：计算机23-3班", size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, "姓    名：", size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, "学    号：", size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, "完成日期：2026年7月", size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_page_break()

    # Manual table of contents.
    add_paragraph(doc, "目  录", size=18, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    toc = [
        "摘  要",
        "1 绪论",
        "1.1 项目背景",
        "1.2 项目目标",
        "2 需求分析",
        "2.1 用户需求",
        "2.2 功能需求",
        "2.3 非功能需求",
        "3 系统总体设计",
        "3.1 系统架构",
        "3.2 技术选型",
        "3.3 数据库设计",
        "4 系统详细设计与实现",
        "4.1 用户认证与个人中心",
        "4.2 在线景点与酒店搜索",
        "4.3 AI旅游助手",
        "4.4 游记社区",
        "4.5 收藏、历史与行程可视化",
        "5 系统测试",
        "6 总结与展望",
        "参考文献",
    ]
    for item in toc:
        p = add_paragraph(doc, item, size=12)
        p.paragraph_format.left_indent = Cm(0.8 if "." in item else 0)
    doc.add_page_break()

    add_heading(doc, "摘  要", 1)
    add_paragraph(
        doc,
        "TravelMind AI 是一个面向旅游场景的个性化推荐与智能规划平台。系统围绕用户输入当前位置、预算、出行天数和旅游偏好后自动生成目的地推荐、出行方式、预算分析、酒店建议、美食打卡点和每日行程规划等功能展开。项目采用 Spring Boot 3.3 + Java 21 构建后端服务，Vue3 + TypeScript + Vite 构建前端界面，结合 JWT、MyBatis Plus、MySQL、Redis、高德地图 POI 与 AI 模型服务，实现前后端分离的智能旅游应用。系统在设计上遵循简洁可维护原则，本地数据库主要保存用户、收藏、历史、游记、AI会话和行程等业务数据，景点与酒店信息优先通过联网 POI 获取，以提高数据真实性和时效性。",
        first_line=True,
    )
    add_paragraph(doc, "关键词：智能旅游；AI推荐；Spring Boot；Vue3；高德地图；行程规划", bold=True)

    add_heading(doc, "1 绪论", 1)
    add_heading(doc, "1.1 项目背景", 2)
    add_paragraph(
        doc,
        "随着自由行和短途旅行需求增加，用户在出行前通常需要同时查询目的地、交通、酒店、景点、美食、预算和路线等信息。传统方式需要在多个平台之间反复切换，信息分散且规划成本较高。TravelMind AI 将旅游搜索、地图 POI、AI问答、预算估算和行程保存整合在同一个系统中，帮助用户以更低成本获得较完整的旅游方案。",
        first_line=True,
    )
    add_heading(doc, "1.2 项目目标", 2)
    add_paragraph(
        doc,
        "本项目目标是完成一套可运行、可演示、可继续扩展的 AI 个性化旅游推荐与智能规划平台。用户可以注册登录，在首页输入出发地、预算、天数和偏好，系统给出推荐方案；用户也可以在景点、酒店、游记、行程和 AI 工具页面完成进一步查询、收藏、记录和管理。",
        first_line=True,
    )

    add_heading(doc, "2 需求分析", 1)
    add_heading(doc, "2.1 用户需求", 2)
    add_paragraph(
        doc,
        "目标用户主要包括学生、情侣、家庭、亲子出游和普通自由行用户。用户希望快速获得真实可查的景点和酒店信息，能够根据预算自动估算费用，能够保存感兴趣的景点、酒店和游记，并在 AI 对话中持续追问旅游方案细节。",
        first_line=True,
    )
    add_heading(doc, "2.2 功能需求", 2)
    add_table(
        doc,
        ["模块", "主要需求"],
        [
            ["用户模块", "注册、登录、JWT认证、个人中心、头像上传、资料维护、收藏和浏览历史查看。"],
            ["景点模块", "联网搜索景点、默认推荐当季适合出行目的地、查看详情、地图定位、收藏和浏览记录。"],
            ["酒店模块", "按目的地搜索附近酒店、查看详情、展示真实 POI 信息、跳转携程等平台按酒店名称搜索预订。"],
            ["AI助手", "问答式聊天、上下文记忆、旅游推荐、行程规划、预算分析、酒店和景点建议。"],
            ["游记社区", "发布游记、上传多张照片、点赞、收藏、评论、查看详情和删除游记。"],
            ["行程规划", "保存 AI 方案，纵向卡片展示计划，详情页展示结构化行程和横向时间轴流程图。"],
        ],
        [Cm(3.2), Cm(12)],
    )
    add_heading(doc, "2.3 非功能需求", 2)
    add_paragraph(
        doc,
        "系统需要具备较好的可用性、响应速度和可维护性。前端界面采用接近 iOS 的毛玻璃、圆角卡片、深浅色模式和流畅动画，保证用户体验；后端接口遵循 RESTful 风格，统一返回结构和异常处理；敏感配置放在环境变量中，避免上传到 GitHub；项目支持 Docker 和 Nginx 部署。",
        first_line=True,
    )

    add_heading(doc, "3 系统总体设计", 1)
    add_heading(doc, "3.1 系统架构", 2)
    add_paragraph(
        doc,
        "系统采用前后端分离架构。前端负责页面渲染、路由跳转、状态管理、地图展示和用户交互；后端负责认证授权、业务接口、数据库访问、AI服务封装、文件上传和第三方服务调用；MySQL 保存业务数据，Redis 可用于缓存和会话辅助，高德地图提供在线 POI 与地图能力，AI服务支持本地 Ollama 和 DeepSeek API 两种方式。",
        first_line=True,
    )
    add_table(
        doc,
        ["层次", "组成", "说明"],
        [
            ["表现层", "Vue3、Naive UI、UnoCSS、ECharts、高德地图JS SDK", "实现首页、景点、酒店、AI工具、游记、行程、个人中心等页面。"],
            ["接口层", "Spring Boot Controller、JWT过滤", "提供 RESTful API，完成请求校验、认证和响应封装。"],
            ["业务层", "AI服务、POI搜索、用户、收藏、历史、游记、行程逻辑", "负责核心业务处理，保持代码简洁。"],
            ["数据层", "MyBatis Plus、MySQL、Redis、文件上传目录", "保存用户和业务数据，支撑查询、保存和历史记录。"],
            ["外部服务", "高德 POI、DeepSeek/Ollama、携程搜索跳转", "提高旅游信息真实性和 AI 推荐能力。"],
        ],
        [Cm(2.6), Cm(5.4), Cm(7.2)],
    )
    add_heading(doc, "3.2 技术选型", 2)
    add_table(
        doc,
        ["类别", "关键技术"],
        [
            ["后端", "Spring Boot 3.3、Java 21、Spring Security、JWT、MyBatis Plus、MySQL 8、Redis、Maven。"],
            ["前端", "Vue3、TypeScript、Vite、Pinia、Vue Router、Axios、Naive UI、UnoCSS、ECharts。"],
            ["地图与AI", "高德地图 JS SDK、高德 POI 搜索、DeepSeek API、本地 Ollama 模型。"],
            ["部署与工具", "Docker、Nginx、Git、GitHub、Knife4j 接口文档。"],
        ],
        [Cm(3), Cm(12.2)],
    )
    add_heading(doc, "3.3 数据库设计", 2)
    add_paragraph(
        doc,
        "数据库设计遵循“本地只保存必须持久化的数据”的原则。景点和酒店详情主要来自联网 POI，本地重点保存用户、收藏、浏览历史、AI会话、AI消息、游记、评论和行程计划等数据。",
        first_line=True,
    )
    add_table(
        doc,
        ["表名", "用途"],
        [
            ["user", "保存用户账号、邮箱、头像、密码、创建时间等信息。"],
            ["favorite", "保存用户收藏的景点、酒店和游记，支持不同类型统一管理。"],
            ["browse_history", "保存浏览历史，同一对象重复浏览时更新到最前。"],
            ["ai_conversation / ai_message", "保存 AI 会话和上下文消息。"],
            ["travel_note / travel_note_comment", "保存游记、图片、点赞、评论等社区内容。"],
            ["travel_plan", "保存用户创建或 AI 生成的行程计划。"],
        ],
        [Cm(4.3), Cm(10.8)],
    )

    add_heading(doc, "4 系统详细设计与实现", 1)
    add_heading(doc, "4.1 用户认证与个人中心", 2)
    add_paragraph(
        doc,
        "用户模块实现注册、登录、JWT认证和个人中心。后端通过 Spring Security 配置接口权限，登录成功后生成 JWT，前端 Axios 在请求头中携带令牌。个人中心支持头像上传、收藏管理和浏览历史展示，收藏和历史以小方块图片卡片呈现，适合移动端和桌面端浏览。",
        first_line=True,
    )
    add_heading(doc, "4.2 在线景点与酒店搜索", 2)
    add_paragraph(
        doc,
        "景点和酒店模块使用高德联网 POI 作为主要数据源。用户可以按城市或关键词搜索景点、酒店，也可以查看默认推荐。详情页展示名称、地址、经纬度、标签、简介和周边信息。对于 POI 图片缺失的情况，系统使用稳定的分类兜底图，保证页面每次启动后不会出现大面积空白图片。",
        first_line=True,
    )
    add_heading(doc, "4.3 AI旅游助手", 2)
    add_paragraph(
        doc,
        "AI模块封装 AiService，支持问答式旅游顾问、行程规划、预算分析、酒店推荐和景点推荐。前端页面改为聊天模式，回答区域独立滚动，保留上下文记录。AI配置支持本地 Ollama 免费模型，也可以通过环境变量切换为 DeepSeek 的 deepseek-chat 模型。",
        first_line=True,
    )
    add_heading(doc, "4.4 游记社区", 2)
    add_paragraph(
        doc,
        "游记社区支持发布图文内容、上传多张照片、点赞、收藏、评论和删除。用户可以点击游记卡片进入详情页，下方展示评论列表并支持发送评论。后端统一 RESTful 路径，解决了删除和评论请求路径不一致导致的接口异常问题。",
        first_line=True,
    )
    add_heading(doc, "4.5 收藏、历史与行程可视化", 2)
    add_paragraph(
        doc,
        "收藏模块支持收藏景点、酒店和游记，并在个人中心统一展示。浏览历史按照业务类型和目标 ID 去重，重复查看时更新到最前。行程模块将保存的方案显示为纵向卡片，详情页按照天数生成横向流程图，节点包括交通、景点、美食和住宿，点击节点可查看 AI 原始建议、避坑贴士和预约方式。",
        first_line=True,
    )

    add_heading(doc, "5 系统测试", 1)
    add_table(
        doc,
        ["测试项", "测试内容", "结果"],
        [
            ["注册登录", "测试注册、登录、JWT失效和未登录访问受限接口。", "通过"],
            ["景点酒店", "测试默认推荐、关键词搜索、详情跳转、地图展示和图片兜底。", "通过"],
            ["AI助手", "测试连续对话、上下文保留、预算分析和模型切换错误提示。", "通过"],
            ["收藏历史", "测试收藏后个人中心展示、重复浏览去重和更新到最前。", "通过"],
            ["游记社区", "测试多图上传、评论发送、详情查看和删除游记。", "通过"],
            ["行程计划", "测试行程保存、卡片列表、结构化详情和流程图弹窗。", "通过"],
        ],
        [Cm(3), Cm(9), Cm(2)],
    )
    add_paragraph(
        doc,
        "通过测试可以看出，系统主要功能能够形成闭环，前后端接口匹配，页面交互基本稳定。由于景点、酒店和部分社交内容依赖外部联网服务，后续仍需加强接口异常兜底、缓存策略和数据质量校验。",
        first_line=True,
    )

    add_heading(doc, "6 总结与展望", 1)
    add_paragraph(
        doc,
        "本次专业实践完成了 TravelMind AI 智能旅游规划平台的设计与开发。项目从需求分析、数据库设计、后端接口、前端页面、AI接入、地图搜索、文件上传到 Git 管理均完成了实践。开发过程中重点解决了图片缺失、在线 POI 搜索、AI模型配置、收藏历史同步、游记评论异常和行程可视化等问题。",
        first_line=True,
    )
    add_paragraph(
        doc,
        "后续可继续优化三个方向：第一，接入更稳定的旅游内容源，提高网红打卡点、美食和酒店推荐的实时性；第二，引入向量检索或知识库，让 AI 对旅游信息的理解更全面；第三，完善部署、监控和移动端适配，使系统更接近真实生产环境。",
        first_line=True,
    )

    add_heading(doc, "参考文献", 1)
    refs = [
        "Spring Boot 官方文档。",
        "Vue.js 官方文档。",
        "MyBatis Plus 官方文档。",
        "高德开放平台开发文档。",
        "DeepSeek API 文档。",
    ]
    for ref in refs:
        add_paragraph(doc, ref, size=12)

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
